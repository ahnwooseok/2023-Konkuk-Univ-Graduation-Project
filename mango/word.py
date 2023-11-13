# import docx

# doc = docx.Document("zzzz.docx")
# # paragraph 정보를 저장할 fullText 리스트 만들기

# fullText = []


# # paragraph(문단)의 각각의 para(줄) 정보를 하나씩 읽고 fullText 리스트에 저장하기

# for para in doc.paragraphs:
#     fullText.append(para.text)

# # 워드 문서 전체 출력하기 (모든 para 값들이 하나의 리스트로 출력됨)

# print(fullText)


# # 워드 문서 전체 출력하기 (word 문서 상의 형식대로 paragraph 를 한줄씩 띄어서 출력하기)

# for i in fullText:
#     print(i)


def getText(filename):
    import docx

    doc = docx.Document(filename)

    fullText = []

    for para in doc.paragraphs:
        fullText.append(para.text)

    return "\n".join(fullText)


print(getText("zzzz.docx"))
